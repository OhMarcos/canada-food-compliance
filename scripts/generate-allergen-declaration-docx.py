"""
Generate Supplier Allergen Declaration as a Word (.docx) file.
Based on Canada's FDR B.01.010.1 - B.01.010.4
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "templates")
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "Supplier_Allergen_Declaration.docx")


def set_cell_shading(cell, color_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elm = tc_pr.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        },
    )
    tc_pr.append(shading_elm)


def add_table_row(table, cells_text, bold=False, shading=None):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "Malgun Gothic"
        run.bold = bold
        if shading:
            set_cell_shading(cell, shading)
    return row


def style_header_row(row, col_count, color_hex="1F4E79"):
    for i in range(col_count):
        cell = row.cells[i]
        set_cell_shading(cell, color_hex)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)


def create_info_table(doc, rows_data):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows_data:
        row = table.add_row()
        c0 = row.cells[0]
        c0.text = ""
        run = c0.paragraphs[0].add_run(label)
        run.font.size = Pt(9)
        run.font.name = "Malgun Gothic"
        run.bold = True
        set_cell_shading(c0, "D6E4F0")
        c0.width = Cm(6)

        c1 = row.cells[1]
        c1.text = value
        for p in c1.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.name = "Malgun Gothic"
        c1.width = Cm(10)
    return table


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # --- Title ---
    title = doc.add_heading("SUPPLIER ALLERGEN DECLARATION", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("공급자 알레르기 유발 물질 선언서")
    run.font.size = Pt(14)
    run.font.name = "Malgun Gothic"
    run.bold = True

    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ref.add_run(
        "Based on Canada's Food and Drug Regulations (FDR) B.01.010.1 - B.01.010.4\n"
        "캐나다 식품의약품규정(FDR) B.01.010.1 - B.01.010.4 기준"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "Malgun Gothic"

    doc.add_paragraph()

    # --- Supplier Information ---
    doc.add_heading("SUPPLIER INFORMATION / 공급자 정보", level=1)
    create_info_table(doc, [
        ("Company Name / 회사명", ""),
        ("Address / 주소", ""),
        ("Contact Person / 담당자", ""),
        ("Phone / 전화번호", ""),
        ("Email / 이메일", ""),
    ])

    doc.add_paragraph()

    # --- Product Information ---
    doc.add_heading("PRODUCT INFORMATION / 제품 정보", level=1)
    create_info_table(doc, [
        ("Product Name / 제품명", ""),
        ("Product Code / 제품 코드", ""),
        ("LOT Number / LOT 번호", ""),
        ("Manufacturing Date / 제조일자", ""),
        ("Specification Version / 사양서 버전", ""),
    ])

    doc.add_paragraph()

    # --- Section A: Allergen Status ---
    doc.add_heading("SECTION A: ALLERGEN STATUS / A항: 알레르겐 상태", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        "Canada's 11 Priority Allergens / 캐나다 11대 우선 알레르기 유발 물질\n"
        "For each allergen, mark (O) in the applicable column. / "
        "각 알레르겐에 대해 해당하는 열에 (O) 표시하세요."
    )
    run.font.size = Pt(9)
    run.font.name = "Malgun Gothic"

    allergen_table = doc.add_table(rows=1, cols=6)
    allergen_table.style = "Table Grid"
    allergen_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "#",
        "Allergen (EN)",
        "알레르겐 (KO)",
        "Intentionally Present\n의도적 포함",
        "Not Present\n미포함",
        "Cross-contact Risk\n교차오염 위험",
    ]
    for i, h in enumerate(headers):
        cell = allergen_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.name = "Malgun Gothic"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_header_row(allergen_table.rows[0], 6)

    allergens = [
        ("1", "Eggs", "계란"),
        ("2", "Milk", "우유"),
        ("3", "Mustard", "겨자"),
        ("4", "Peanuts", "땅콩"),
        ("5", "Crustaceans & Molluscs", "갑각류 & 연체동물"),
        ("6", "Fish", "생선"),
        ("7", "Sesame Seeds", "참깨"),
        ("8", "Soy", "대두"),
        ("9", "Sulphites (>=10ppm)", "아황산염 (>=10ppm)"),
        ("10", "Tree Nuts", "견과류"),
        ("11", "Wheat & Triticale", "밀 & 트리티케일"),
    ]
    for num, en, ko in allergens:
        add_table_row(allergen_table, [num, en, ko, "", "", ""])

    p = doc.add_paragraph()
    run = p.add_run(
        "\nIf Tree Nuts is marked, specify type(s) / 견과류 해당 시 종류 명시: "
        "_______________________________________________"
    )
    run.font.size = Pt(9)
    run.font.name = "Malgun Gothic"

    doc.add_paragraph()

    # --- Section B: Ingredient List ---
    doc.add_heading(
        "SECTION B: COMPLETE INGREDIENT & SUB-INGREDIENT LIST / B항: 전체 성분 및 부원료 목록",
        level=1,
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "List ALL ingredients including sub-ingredients, processing aids, and carriers.\n"
        "모든 성분(부원료, 가공보조제, 캐리어 포함)을 나열하세요."
    )
    run.font.size = Pt(9)
    run.font.name = "Malgun Gothic"

    ing_table = doc.add_table(rows=1, cols=5)
    ing_table.style = "Table Grid"
    ing_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    ing_headers = [
        "#",
        "Ingredient / 성분",
        "Function / 기능",
        "Contains Allergen?\n알레르겐 포함?",
        "If Yes, Which?\n해당 시 종류",
    ]
    for i, h in enumerate(ing_headers):
        cell = ing_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.name = "Malgun Gothic"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_header_row(ing_table.rows[0], 5)

    for n in range(1, 11):
        add_table_row(ing_table, [str(n), "", "", "Yes / No", ""])

    p = doc.add_paragraph()
    run = p.add_run("(Add rows as needed / 필요 시 행을 추가하세요)")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "Malgun Gothic"

    doc.add_paragraph()

    # --- Section C: Cross-Contact ---
    doc.add_heading(
        "SECTION C: CROSS-CONTACT DETAILS / C항: 교차오염 상세", level=1
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "C-1. Does the manufacturing facility process any of Canada's 11 priority "
        "allergens on shared equipment or in the same facility?\n"
        "제조 시설에서 캐나다 11대 우선 알레르겐을 공유 장비 또는 동일 시설에서 취급합니까?"
    )
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.name = "Malgun Gothic"

    p = doc.add_paragraph()
    run = p.add_run("      \u2610  YES / 예          \u2610  NO / 아니오")
    run.font.size = Pt(10)
    run.font.name = "Malgun Gothic"

    p = doc.add_paragraph()
    run = p.add_run(
        "C-2. If YES, complete the following / 해당 시 아래를 작성하세요:"
    )
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.name = "Malgun Gothic"

    cross_table = doc.add_table(rows=1, cols=4)
    cross_table.style = "Table Grid"
    cross_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cross_headers = [
        "Allergen on Shared Line\n공유 라인의 알레르겐",
        "Equipment Shared\n공유 장비",
        "Cleaning Method\n세척 방법",
        "Cleaning Validated?\n세척 검증 여부",
    ]
    for i, h in enumerate(cross_headers):
        cell = cross_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.name = "Malgun Gothic"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_header_row(cross_table.rows[0], 4)

    for _ in range(4):
        add_table_row(cross_table, ["", "", "", "Yes / No"])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(
        "C-3. Is the facility dedicated allergen-free for any specific allergen?\n"
        "특정 알레르겐에 대해 전용(무알레르겐) 시설입니까?"
    )
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.name = "Malgun Gothic"

    p = doc.add_paragraph()
    run = p.add_run(
        "      \u2610  YES / 예  (specify / 명시: ___________________________)          "
        "\u2610  NO / 아니오"
    )
    run.font.size = Pt(10)
    run.font.name = "Malgun Gothic"

    doc.add_paragraph()

    # --- Section D: Special Verification ---
    doc.add_heading(
        "SECTION D: SPECIAL VERIFICATION / D항: 특별 확인 사항", level=1
    )

    p = doc.add_paragraph()
    run = p.add_run("For Gochujang / Fermented Products / 고추장 및 발효 제품용")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.name = "Malgun Gothic"

    goch_table = doc.add_table(rows=1, cols=2)
    goch_table.style = "Table Grid"
    goch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Question / 질문", "Answer / 답변"]):
        cell = goch_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.name = "Malgun Gothic"
    style_header_row(goch_table.rows[0], 2)

    goch_questions = [
        "Source of Malt (맥아 원료): Barley (보리) or Wheat (밀)?",
        "If Barley Malt, does it contain Gluten? / 보리 맥아인 경우 글루텐 포함?",
        "Does fermentation process use any allergen-containing starter?\n발효 과정에서 알레르겐 포함 종균 사용?",
    ]
    for q in goch_questions:
        add_table_row(goch_table, [q, ""])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("For Soy-based Ingredients / 대두 기반 원료용")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.name = "Malgun Gothic"

    soy_table = doc.add_table(rows=1, cols=2)
    soy_table.style = "Table Grid"
    soy_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Question / 질문", "Answer / 답변"]):
        cell = soy_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.name = "Malgun Gothic"
    style_header_row(soy_table.rows[0], 2)

    soy_questions = [
        "Is the soy non-GMO? / 대두가 Non-GMO입니까?",
        "IP (Identity Preserved) documentation available? / IP 문서 제공 가능?",
    ]
    for q in soy_questions:
        add_table_row(soy_table, [q, ""])

    doc.add_paragraph()

    # --- Section E: Change Notification ---
    doc.add_heading(
        "SECTION E: CHANGE NOTIFICATION / E항: 변경 통보", level=1
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "The supplier agrees to notify the buyer IN WRITING at least 30 days "
        "prior to any of the following changes:\n"
        "공급자는 아래 변경 사항 발생 시 최소 30일 전에 서면으로 구매자에게 "
        "통보할 것에 동의합니다:"
    )
    run.font.size = Pt(9)
    run.font.name = "Malgun Gothic"

    changes = [
        "Change in ingredient formulation / 배합 변경",
        "Change in sub-ingredient supplier / 부원료 공급자 변경",
        "Change in manufacturing facility / 제조 시설 변경",
        "Change in allergen status (including cross-contact) / 알레르겐 상태 변경 (교차오염 포함)",
        "Change in cleaning or sanitation procedures / 세척 또는 위생 절차 변경",
    ]
    for c in changes:
        p = doc.add_paragraph()
        run = p.add_run(f"  \u2610  {c}")
        run.font.size = Pt(9)
        run.font.name = "Malgun Gothic"

    doc.add_paragraph()

    # --- Section F: Certification ---
    doc.add_heading("SECTION F: CERTIFICATION / F항: 확인 서명", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "I certify that the information provided in this declaration is accurate, "
        "complete, and reflects the current status of the product identified above. "
        "I understand that this declaration will be used by the buyer to meet "
        "Canadian food safety requirements under the Safe Food for Canadians "
        "Regulations (SFCR) and Food and Drug Regulations (FDR). I will notify "
        "the buyer immediately of any changes that affect the allergen status "
        "of this product."
    )
    run.font.size = Pt(9)
    run.font.name = "Malgun Gothic"

    p = doc.add_paragraph()
    run = p.add_run(
        "본인은 이 선언서에 제공된 정보가 정확하고 완전하며, 위에 명시된 제품의 "
        "현재 상태를 반영함을 확인합니다. 본 선언서가 캐나다 식품안전규정(SFCR) "
        "및 식품의약품규정(FDR)에 따른 구매자의 캐나다 식품 안전 요건 충족을 위해 "
        "사용됨을 이해합니다. 본 제품의 알레르겐 상태에 영향을 미치는 변경 사항이 "
        "있을 경우 즉시 구매자에게 통보하겠습니다."
    )
    run.font.size = Pt(9)
    run.font.name = "Malgun Gothic"

    doc.add_paragraph()

    create_info_table(doc, [
        ("Name / 성명", ""),
        ("Title / 직위", ""),
        ("Signature / 서명", ""),
        ("Date / 날짜", ""),
        ("Validity Period / 유효 기간", ""),
    ])

    doc.add_paragraph()

    # --- Footer ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This document should be updated per LOT or upon any formulation change.\n"
        "이 문서는 매 LOT마다 또는 배합 변경 시 갱신해야 합니다."
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "Malgun Gothic"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Legal References: FDR B.01.010.1-B.01.010.4 | SFCR Part 4, s.52-58 | "
        "CFIA Allergen Labelling Guidelines"
    )
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = "Malgun Gothic"

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
