"""
Generate Supplier Letter of Guarantee (LoG) as a Word (.docx) file.
For Canadian food import compliance under SFCR Part 4.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "templates")
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "Supplier_Letter_of_Guarantee.docx")

FONT = "Malgun Gothic"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elm = tc_pr.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        },
    )
    tc_pr.append(shading_elm)


def styled_run(paragraph, text, size=9, bold=False, italic=False, color=None):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.name = FONT
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def style_header_row(row, col_count, color_hex="1F4E79"):
    for i in range(col_count):
        cell = row.cells[i]
        set_cell_shading(cell, color_hex)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)


def add_header_table(doc, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.name = FONT
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_header_row(table.rows[0], len(headers))
    return table


def add_table_row(table, cells_text):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(9)
        run.font.name = FONT
    return row


def create_info_table(doc, rows_data):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows_data:
        row = table.add_row()
        c0 = row.cells[0]
        c0.text = ""
        run = c0.paragraphs[0].add_run(label)
        run.font.size = Pt(9)
        run.font.name = FONT
        run.bold = True
        set_cell_shading(c0, "D6E4F0")
        c0.width = Cm(6)

        c1 = row.cells[1]
        c1.text = value
        for p in c1.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.name = FONT
        c1.width = Cm(10)
    return table


def add_checkbox_line(doc, text, indent="  "):
    p = doc.add_paragraph()
    styled_run(p, f"{indent}\u2610  {text}")
    return p


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ================================================================
    # TITLE
    # ================================================================
    title = doc.add_heading("LETTER OF GUARANTEE", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "공급자 보증서", size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(
        p,
        "For Canadian Food Import Compliance\n"
        "캐나다 식품 수입 컴플라이언스용\n\n"
        "Legal Basis: SFCR Part 4 s.52-58 (Preventive Control Plans)\n"
        "법적 근거: SFCR 제4부 제52-58조 (예방관리계획)",
        size=8,
        color=(0x66, 0x66, 0x66),
    )

    doc.add_paragraph()

    # ================================================================
    # DOCUMENT INFORMATION
    # ================================================================
    doc.add_heading("DOCUMENT INFORMATION / 문서 정보", level=1)
    create_info_table(doc, [
        ("Document No. / 문서 번호", ""),
        ("Issue Date / 발행일", ""),
        ("Revision No. / 개정 번호", ""),
        ("Valid Until / 유효 기한", ""),
    ])

    doc.add_paragraph()

    # ================================================================
    # SUPPLIER INFORMATION
    # ================================================================
    doc.add_heading("SUPPLIER INFORMATION / 공급자 정보", level=1)
    create_info_table(doc, [
        ("Company Name / 회사명", ""),
        ("Business Registration No. / 사업자등록번호", ""),
        ("Address / 주소", ""),
        ("Manufacturing Facility Address / 제조 시설 주소", ""),
        ("Contact Person / 담당자", ""),
        ("Title / 직위", ""),
        ("Phone / 전화번호", ""),
        ("Email / 이메일", ""),
    ])

    doc.add_paragraph()

    # ================================================================
    # BUYER INFORMATION
    # ================================================================
    doc.add_heading("BUYER INFORMATION / 구매자 정보", level=1)
    create_info_table(doc, [
        ("Company Name / 회사명", ""),
        ("SFC Licence No. / SFC 라이선스 번호", ""),
        ("Address / 주소", ""),
        ("Contact Person / 담당자", ""),
    ])

    doc.add_paragraph()

    # ================================================================
    # PRODUCT INFORMATION
    # ================================================================
    doc.add_heading("PRODUCT(S) COVERED / 대상 제품", level=1)

    prod_table = add_header_table(doc, [
        "#",
        "Product Name\n제품명",
        "Product Code\n제품 코드",
        "HS Code\nHS 코드",
        "Specification Ver.\n사양서 버전",
    ])
    for n in range(1, 6):
        add_table_row(prod_table, [str(n), "", "", "", ""])

    doc.add_paragraph()

    # ================================================================
    # SECTION 1: FOOD SAFETY GUARANTEE
    # ================================================================
    doc.add_heading(
        "SECTION 1: FOOD SAFETY GUARANTEE / 1항: 식품 안전 보증", level=1
    )

    p = doc.add_paragraph()
    styled_run(
        p,
        "The supplier guarantees that all products listed above: / "
        "공급자는 위에 나열된 모든 제품이 다음을 충족함을 보증합니다:",
        bold=True,
    )

    guarantees_1 = [
        (
            "Are manufactured in a facility that meets Good Manufacturing Practices (GMP) "
            "and maintains a HACCP-based food safety system.\n"
            "우수제조관리기준(GMP)을 충족하고 HACCP 기반 식품 안전 시스템을 유지하는 시설에서 "
            "제조됩니다."
        ),
        (
            "Are safe and suitable for human consumption.\n"
            "안전하며 인간이 섭취하기에 적합합니다."
        ),
        (
            "Are manufactured under sanitary conditions with adequate pest control, "
            "employee hygiene, and facility maintenance programs.\n"
            "적절한 해충 관리, 직원 위생 및 시설 유지 관리 프로그램 하에 위생적인 조건에서 "
            "제조됩니다."
        ),
        (
            "Comply with all applicable Korean food safety laws and regulations "
            "(식품위생법, 식품의 기준 및 규격).\n"
            "한국의 모든 관련 식품 안전 법규(식품위생법, 식품의 기준 및 규격)를 준수합니다."
        ),
    ]

    for g in guarantees_1:
        add_checkbox_line(doc, g)

    doc.add_paragraph()

    # ================================================================
    # SECTION 2: INGREDIENT & ADDITIVE COMPLIANCE
    # ================================================================
    doc.add_heading(
        "SECTION 2: INGREDIENT & ADDITIVE COMPLIANCE / 2항: 성분 및 첨가물 적합성",
        level=1,
    )

    p = doc.add_paragraph()
    styled_run(p, "The supplier guarantees that: / 공급자는 다음을 보증합니다:", bold=True)

    guarantees_2 = [
        (
            "A complete and accurate ingredient list (including all sub-ingredients, "
            "processing aids, and carriers) has been provided to the buyer.\n"
            "완전하고 정확한 성분 목록(모든 부원료, 가공보조제, 캐리어 포함)을 구매자에게 "
            "제공하였습니다."
        ),
        (
            "All food additives used in the product(s) are permitted under Canada's "
            "Food and Drug Regulations (FDR) and Health Canada's Lists of Permitted "
            "Food Additives.\n"
            "제품에 사용된 모든 식품 첨가물은 캐나다 식품의약품규정(FDR) 및 Health Canada의 "
            "허용 식품 첨가물 목록에서 허용됩니다."
        ),
        (
            "Food additives are used within the maximum levels of use prescribed "
            "by Canadian regulations.\n"
            "식품 첨가물은 캐나다 규정에서 정한 최대 사용량 이내로 사용됩니다."
        ),
        (
            "No additives, preservatives, or processing aids prohibited in Canada "
            "are present in the product(s).\n"
            "캐나다에서 금지된 첨가물, 보존료 또는 가공보조제가 제품에 포함되어 있지 않습니다."
        ),
    ]

    for g in guarantees_2:
        add_checkbox_line(doc, g)

    doc.add_paragraph()

    # ================================================================
    # SECTION 3: ALLERGEN GUARANTEE
    # ================================================================
    doc.add_heading(
        "SECTION 3: ALLERGEN GUARANTEE / 3항: 알레르겐 보증", level=1
    )

    p = doc.add_paragraph()
    styled_run(
        p,
        "The supplier guarantees the following regarding Canada's 11 Priority Allergens "
        "(FDR B.01.010.1):\n"
        "공급자는 캐나다 11대 우선 알레르기 유발 물질(FDR B.01.010.1)에 대해 다음을 보증합니다:",
        bold=True,
    )

    guarantees_3 = [
        (
            "All priority allergens intentionally present in the product(s) have been "
            "accurately disclosed to the buyer.\n"
            "제품에 의도적으로 포함된 모든 우선 알레르겐을 구매자에게 정확히 공개하였습니다."
        ),
        (
            "All potential allergen cross-contact risks from shared equipment or "
            "facility have been disclosed.\n"
            "공유 장비 또는 시설로 인한 모든 잠재적 알레르겐 교차오염 위험을 공개하였습니다."
        ),
        (
            "An allergen control program is in place at the manufacturing facility.\n"
            "제조 시설에 알레르겐 관리 프로그램이 시행되고 있습니다."
        ),
    ]

    for g in guarantees_3:
        add_checkbox_line(doc, g)

    p = doc.add_paragraph()
    styled_run(
        p,
        "Allergens present in product(s) / 제품에 포함된 알레르겐:",
        bold=True,
    )

    allergen_table = add_header_table(doc, [
        "Allergen / 알레르겐",
        "Present\n포함",
        "Not Present\n미포함",
        "Cross-contact\n교차오염",
    ])

    allergens = [
        "Eggs / 계란", "Milk / 우유", "Mustard / 겨자", "Peanuts / 땅콩",
        "Crustaceans & Molluscs / 갑각류 & 연체동물", "Fish / 생선",
        "Sesame Seeds / 참깨", "Soy / 대두",
        "Sulphites (>=10ppm) / 아황산염", "Tree Nuts / 견과류",
        "Wheat & Triticale / 밀 & 트리티케일",
    ]
    for a in allergens:
        add_table_row(allergen_table, [a, "", "", ""])

    doc.add_paragraph()

    # ================================================================
    # SECTION 4: CONTAMINANT COMPLIANCE
    # ================================================================
    doc.add_heading(
        "SECTION 4: CONTAMINANT COMPLIANCE / 4항: 오염물질 적합성", level=1
    )

    p = doc.add_paragraph()
    styled_run(p, "The supplier guarantees that: / 공급자는 다음을 보증합니다:", bold=True)

    guarantees_4 = [
        (
            "Pesticide residues in the product(s) comply with Canada's Maximum Residue "
            "Limits (MRLs) as set out in the Pest Control Products Act.\n"
            "제품의 잔류농약은 캐나다 해충관리제품법에서 정한 최대잔류허용기준(MRL)을 준수합니다."
        ),
        (
            "Heavy metal levels (lead, cadmium, arsenic, mercury) comply with Canadian "
            "regulatory limits.\n"
            "중금속 수치(납, 카드뮴, 비소, 수은)는 캐나다 규제 한도를 준수합니다."
        ),
        (
            "Mycotoxin levels (where applicable) comply with Canadian limits.\n"
            "곰팡이 독소 수치(해당 시)는 캐나다 한도를 준수합니다."
        ),
        (
            "The product(s) are free from foreign objects and physical contaminants.\n"
            "제품에 이물질 및 물리적 오염물질이 없습니다."
        ),
    ]

    for g in guarantees_4:
        add_checkbox_line(doc, g)

    doc.add_paragraph()

    # ================================================================
    # SECTION 5: NON-GMO STATUS (IF APPLICABLE)
    # ================================================================
    doc.add_heading(
        "SECTION 5: NON-GMO STATUS (IF APPLICABLE) / 5항: Non-GMO 상태 (해당 시)",
        level=1,
    )

    p = doc.add_paragraph()
    styled_run(p, "Does this section apply? / 이 항목이 해당됩니까?")

    p = doc.add_paragraph()
    styled_run(p, "  \u2610  YES / 예 (complete below / 아래 작성)          \u2610  NO / 아니오 (skip to Section 6 / 6항으로 이동)", size=10)

    p = doc.add_paragraph()
    styled_run(
        p,
        "If applicable, the supplier guarantees that: / 해당 시, 공급자는 다음을 보증합니다:",
        bold=True,
    )

    guarantees_5 = [
        (
            "The product(s) are derived from non-genetically modified organisms.\n"
            "제품은 유전자 변형 생물체(GMO)에서 유래하지 않았습니다."
        ),
        (
            "Identity Preserved (IP) or equivalent segregation systems are maintained "
            "throughout the supply chain.\n"
            "공급망 전체에 걸쳐 IP(Identity Preserved) 또는 동등한 분리 시스템이 유지됩니다."
        ),
        (
            "Supporting documentation (CoA with non-GMO test results, IP certificates) "
            "is available upon request.\n"
            "뒷받침 문서(Non-GMO 검사 결과가 포함된 CoA, IP 인증서)를 요청 시 제공할 수 있습니다."
        ),
    ]

    for g in guarantees_5:
        add_checkbox_line(doc, g)

    gmo_table = add_header_table(doc, [
        "Ingredient / 원료",
        "GMO Risk Level\nGMO 위험도",
        "Non-GMO Verified?\n검증 여부",
        "Test Method\n검사 방법",
    ])
    gmo_rows = [
        ("Soy Protein Concentrate / 대두 단백 농축물", "HIGH / 높음", "", ""),
        ("Soybean Oil / 대두유", "HIGH / 높음", "", ""),
        ("Cornstarch / 옥수수 전분", "HIGH / 높음", "", ""),
        ("Rice Syrup / 쌀 시럽", "LOW / 낮음", "", ""),
        ("Cane Sugar / 사탕수수 설탕", "LOW / 낮음", "", ""),
    ]
    for row_data in gmo_rows:
        add_table_row(gmo_table, list(row_data))

    doc.add_paragraph()

    # ================================================================
    # SECTION 6: TRACEABILITY
    # ================================================================
    doc.add_heading(
        "SECTION 6: TRACEABILITY / 6항: 추적성", level=1
    )

    p = doc.add_paragraph()
    styled_run(p, "The supplier guarantees that: / 공급자는 다음을 보증합니다:", bold=True)

    guarantees_6 = [
        (
            "A traceability system is in place that can identify the source and "
            "destination of all ingredients and finished products.\n"
            "모든 원재료 및 완제품의 출처와 목적지를 식별할 수 있는 추적 시스템이 "
            "구축되어 있습니다."
        ),
        (
            "Lot coding is applied to all products supplied.\n"
            "공급되는 모든 제품에 로트 코딩이 적용됩니다."
        ),
        (
            "Traceability records are retained for a minimum of 2 years.\n"
            "추적 기록은 최소 2년간 보관됩니다."
        ),
        (
            "In the event of a recall or safety issue, the supplier will cooperate "
            "fully and provide traceability information within 24 hours of request.\n"
            "리콜 또는 안전 문제 발생 시, 공급자는 전적으로 협력하며 요청 후 24시간 이내에 "
            "추적 정보를 제공합니다."
        ),
    ]

    for g in guarantees_6:
        add_checkbox_line(doc, g)

    doc.add_paragraph()

    # ================================================================
    # SECTION 7: CHANGE NOTIFICATION
    # ================================================================
    doc.add_heading(
        "SECTION 7: CHANGE NOTIFICATION / 7항: 변경 통보", level=1
    )

    p = doc.add_paragraph()
    styled_run(
        p,
        "The supplier agrees to notify the buyer IN WRITING at least 30 days "
        "prior to any of the following changes:\n"
        "공급자는 아래 변경 사항 발생 시 최소 30일 전에 서면으로 구매자에게 "
        "통보할 것에 동의합니다:",
    )

    changes = [
        "Change in ingredient formulation or sub-ingredients / 배합 또는 부원료 변경",
        "Change in manufacturing facility or production line / 제조 시설 또는 생산 라인 변경",
        "Change in allergen status (including cross-contact) / 알레르겐 상태 변경 (교차오염 포함)",
        "Change in food additive usage / 식품 첨가물 사용 변경",
        "Change in cleaning or sanitation procedures / 세척 또는 위생 절차 변경",
        "Change in GMO status or supplier of raw materials / GMO 상태 또는 원재료 공급자 변경",
        "Any regulatory non-compliance or recall at the supplier facility / 공급자 시설의 규제 부적합 또는 리콜",
    ]

    for c in changes:
        add_checkbox_line(doc, c)

    doc.add_paragraph()

    # ================================================================
    # SECTION 8: SUPPORTING DOCUMENTS
    # ================================================================
    doc.add_heading(
        "SECTION 8: SUPPORTING DOCUMENTS PROVIDED / 8항: 제출 첨부 서류", level=1
    )

    p = doc.add_paragraph()
    styled_run(
        p,
        "Check all documents attached with this Letter of Guarantee:\n"
        "본 보증서와 함께 첨부된 서류를 모두 체크하세요:",
        bold=True,
    )

    docs_list = [
        "Certificate of Analysis (CoA) / 성적분석서",
        "Complete Specification Sheet / 완전한 사양서",
        "Allergen Declaration / 알레르기 유발 물질 선언서",
        "Non-GMO Certificate or CoA / Non-GMO 인증서 또는 성적서",
        "IP (Identity Preserved) Documentation / IP 문서",
        "HACCP Certificate / HACCP 인증서",
        "Food Safety Audit Report / 식품 안전 감사 보고서",
        "Phytosanitary Certificate / 식물검역증명서",
        "CKFTA Certificate of Origin / 한-캐 FTA 원산지 증명서",
        "Other (specify) / 기타 (명시): ___________________________",
    ]

    for d in docs_list:
        add_checkbox_line(doc, d)

    doc.add_paragraph()

    # ================================================================
    # CERTIFICATION & SIGNATURE
    # ================================================================
    doc.add_heading("CERTIFICATION & SIGNATURE / 확인 및 서명", level=1)

    p = doc.add_paragraph()
    styled_run(
        p,
        "I, the undersigned, as an authorized representative of the supplier company "
        "identified above, hereby certify and guarantee that all information provided "
        "in this Letter of Guarantee is true, accurate, and complete to the best of "
        "my knowledge. I understand that this document will be used by the buyer as "
        "part of their Preventive Control Plan (PCP) under Canada's Safe Food for "
        "Canadians Regulations (SFCR) and Food and Drug Regulations (FDR).",
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    styled_run(
        p,
        "본인은 위에 명시된 공급자 회사의 권한 있는 대표자로서, 본 보증서에 제공된 모든 "
        "정보가 본인이 아는 한 사실이며, 정확하고 완전함을 확인하고 보증합니다. 본 문서가 "
        "캐나다 식품안전규정(SFCR) 및 식품의약품규정(FDR)에 따른 구매자의 예방관리계획(PCP)의 "
        "일부로 사용됨을 이해합니다.",
    )

    doc.add_paragraph()

    # Supplier signature
    p = doc.add_paragraph()
    styled_run(p, "SUPPLIER / 공급자", size=11, bold=True)

    create_info_table(doc, [
        ("Name / 성명", ""),
        ("Title / 직위", ""),
        ("Signature / 서명", ""),
        ("Date / 날짜", ""),
        ("Company Stamp / 회사 직인", ""),
    ])

    doc.add_paragraph()

    # Buyer acknowledgment
    p = doc.add_paragraph()
    styled_run(p, "BUYER ACKNOWLEDGMENT / 구매자 확인", size=11, bold=True)

    create_info_table(doc, [
        ("Received By / 접수자", ""),
        ("Title / 직위", ""),
        ("Date Received / 접수일", ""),
        ("Review Status / 검토 상태", "Approved / Pending / Rejected"),
        ("Notes / 비고", ""),
    ])

    doc.add_paragraph()

    # ================================================================
    # FOOTER
    # ================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(
        p,
        "This Letter of Guarantee is valid for the period specified above.\n"
        "It must be renewed annually or upon any material change.\n"
        "본 보증서는 위에 명시된 기간 동안 유효합니다.\n"
        "매년 또는 중요한 변경 시 갱신해야 합니다.",
        size=8,
        italic=True,
        color=(0x66, 0x66, 0x66),
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(
        p,
        "Legal References: SFCR Part 4 s.52-58 | SFCR Part 5 s.59-72 | "
        "FDR B.01.010.1-B.01.010.4 | FDR B.01.301-B.01.305",
        size=7,
        color=(0x99, 0x99, 0x99),
    )

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
